import html
import io
import logging
import re
import uuid
from collections.abc import AsyncGenerator
from enum import Enum
from typing import IO, Optional, cast

import pymupdf
from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import (
    AnalyzeDocumentRequest,
    AnalyzeResult,
    DocumentFigure,
    DocumentTable,
)
from azure.core.credentials import AzureKeyCredential
from azure.core.credentials_async import AsyncTokenCredential
from azure.core.exceptions import HttpResponseError
from PIL import Image
from pypdf import PdfReader

from .page import ImageOnPage, Page
from .parser import Parser

logger = logging.getLogger("scripts")


def _build_heading_map_from_docx(content_bytes: bytes) -> dict[str, int]:
    """Return {heading_text: correct_level} by reading the real Word paragraph styles.

    python-docx reads the actual XML style names (Heading 1, Heading 2 …), which are
    always correct regardless of how the document was visually formatted.  Only
    headings whose text maps to exactly one level are included; ambiguous entries are
    dropped so we never make a wrong correction.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(io.BytesIO(content_bytes))
        text_levels: dict[str, set[int]] = {}
        for para in doc.paragraphs:
            style = para.style.name
            if style.startswith("Heading"):
                text = para.text.strip()
                if text:
                    try:
                        level = int(style.split()[-1])
                    except (ValueError, IndexError):
                        continue
                    text_levels.setdefault(text, set()).add(level)
        return {text: next(iter(lvls)) for text, lvls in text_levels.items() if len(lvls) == 1}
    except Exception as exc:
        logger.warning("Could not build heading map from DOCX: %s", exc)
        return {}


def _build_heading_map_from_pdf_toc(content_bytes: bytes) -> dict[str, int]:
    """Return {heading_text: correct_level} from a PDF's embedded table of contents.

    Uses pymupdf's get_toc().  Returns an empty dict if no TOC is present, in which
    case the caller falls back to DI's heading markers unchanged.
    """
    try:
        doc = pymupdf.open(stream=io.BytesIO(content_bytes))
        toc = doc.get_toc()
        doc.close()
        if not toc:
            return {}
        text_levels: dict[str, set[int]] = {}
        for item in toc:
            level = item[0]
            text = item[1].strip()
            if text:
                text_levels.setdefault(text, set()).add(level)
        return {text: next(iter(lvls)) for text, lvls in text_levels.items() if len(lvls) == 1}
    except Exception as exc:
        logger.warning("Could not build heading map from PDF TOC: %s", exc)
        return {}


def _patch_heading_levels(text: str, heading_map: dict[str, int]) -> str:
    """Rewrite DI-assigned # markers in *text* using the correct levels from heading_map."""
    if not heading_map:
        return text

    def replace_heading(m: re.Match) -> str:
        heading_text = m.group(2).strip()
        if heading_text in heading_map:
            return "#" * heading_map[heading_text] + " " + heading_text
        return m.group(0)

    return re.sub(r"^(#{1,6}) +(.+)$", replace_heading, text, flags=re.MULTILINE)


class LocalPdfParser(Parser):
    """Concrete parser backed by PyPDF that can parse PDFs into pages"""

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        logger.info("Extracting text from '%s' using local PDF parser (pypdf)", content.name)

        reader = PdfReader(content)
        pages = reader.pages
        offset = 0
        for page_num, p in enumerate(pages):
            page_text = p.extract_text()
            yield Page(page_num=page_num, offset=offset, text=page_text)
            offset += len(page_text)


class DocumentAnalysisParser(Parser):
    """Concrete parser backed by Azure AI Document Intelligence"""

    def __init__(
        self,
        endpoint: str,
        credential: AsyncTokenCredential | AzureKeyCredential,
        model_id: str = "prebuilt-layout",
        process_figures: bool = False,
    ) -> None:
        self.model_id = model_id
        self.endpoint = endpoint
        self.credential = credential
        self.process_figures = process_figures

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        logger.info("Extracting text from '%s' using Azure Document Intelligence", content.name)

        async with DocumentIntelligenceClient(
            endpoint=self.endpoint, credential=self.credential
        ) as document_intelligence_client:
            try:
                content.seek(0)
            except Exception:
                pass
            content_bytes = content.read()

            # Build a heading-level correction map from the native document parser.
            # Azure DI's markdown heading levels can diverge from the actual Word/PDF
            # heading styles.  For DOCX we use python-docx (reads real XML styles);
            # for PDF we use pymupdf's TOC when available.
            file_name = getattr(content, "name", "") or ""
            heading_map: dict[str, int] = {}
            if file_name.lower().endswith(".docx"):
                heading_map = _build_heading_map_from_docx(content_bytes)
                if heading_map:
                    logger.info(
                        "Built heading correction map from python-docx: %d unambiguous heading(s)",
                        len(heading_map),
                    )
            elif file_name.lower().endswith(".pdf"):
                heading_map = _build_heading_map_from_pdf_toc(content_bytes)
                if heading_map:
                    logger.info(
                        "Built heading correction map from PDF TOC: %d heading(s)",
                        len(heading_map),
                    )

            poller = None
            doc_for_pymupdf = None

            if self.process_figures:
                try:
                    poller = await document_intelligence_client.begin_analyze_document(
                        model_id="prebuilt-layout",
                        body=AnalyzeDocumentRequest(bytes_source=content_bytes),
                        output=["figures"],
                        features=["ocrHighResolution"],
                        output_content_format="markdown",
                    )
                    doc_for_pymupdf = pymupdf.open(stream=io.BytesIO(content_bytes))
                except HttpResponseError as e:
                    logger.error("Error analyzing document for media: %s. Proceeding with standard analysis.", e)
                    poller = None

            if poller is None:
                poller = await document_intelligence_client.begin_analyze_document(
                    model_id=self.model_id,
                    body=AnalyzeDocumentRequest(bytes_source=content_bytes),
                    # Request markdown output so heading levels (# / ## / ###)
                    # are always present for both DOCX and PDF inputs.
                    # For DOCX: DI maps Word heading styles directly to # markers.
                    # For PDF:  DI infers headings from font/size heuristics and
                    #           emits them as # markers only when markdown is requested.
                    output_content_format="markdown",
                )
            analyze_result: AnalyzeResult = await poller.result()

            offset = 0

            for page in analyze_result.pages:
                tables_on_page = [
                    table
                    for table in (analyze_result.tables or [])
                    if table.bounding_regions and table.bounding_regions[0].page_number == page.page_number
                ]
                figures_on_page = []
                if self.process_figures:
                    figures_on_page = [
                        figure
                        for figure in (analyze_result.figures or [])
                        if figure.bounding_regions and figure.bounding_regions[0].page_number == page.page_number
                    ]
                    logger.info(
                        "Page %d: found %d figure(s)",
                        page.page_number,
                        len(figures_on_page),
                    )
                page_images: list[ImageOnPage] = []
                page_tables: list[str] = []

                class ObjectType(Enum):
                    NONE = -1
                    TABLE = 0
                    FIGURE = 1

                MaskEntry = tuple[ObjectType, Optional[int]]

                page_offset = page.spans[0].offset
                page_length = page.spans[0].length
                mask_chars: list[MaskEntry] = cast(list[MaskEntry], [(ObjectType.NONE, None)] * page_length)
                
                for table_idx, table in enumerate(tables_on_page):
                    for span in table.spans:
                        for i in range(span.length):
                            idx = span.offset - page_offset + i
                            if idx >= 0 and idx < page_length:
                                mask_chars[idx] = (ObjectType.TABLE, table_idx)
                
                for figure_idx, figure in enumerate(figures_on_page):
                    for span in figure.spans:
                        for i in range(span.length):
                            idx = span.offset - page_offset + i
                            if idx >= 0 and idx < page_length:
                                mask_chars[idx] = (ObjectType.FIGURE, figure_idx)

                page_text = ""
                added_objects: set[MaskEntry] = set()
                for idx, mask_char in enumerate(mask_chars):
                    object_type, object_idx = mask_char
                    if object_type == ObjectType.NONE:
                        page_text += analyze_result.content[page_offset + idx]
                    elif object_type == ObjectType.TABLE:
                        if mask_char not in added_objects:
                            table_html = DocumentAnalysisParser.table_to_html(tables_on_page[object_idx])
                            page_tables.append(table_html)
                            page_text += table_html
                            added_objects.add(mask_char)
                    elif object_type == ObjectType.FIGURE:
                        if doc_for_pymupdf is None:
                            continue
                        if mask_char not in added_objects:
                            image_on_page = await DocumentAnalysisParser.figure_to_image(
                                doc_for_pymupdf, figures_on_page[object_idx]
                            )
                            page_images.append(image_on_page)
                            page_text += image_on_page.placeholder
                            added_objects.add(mask_char)
                            logger.info(
                                "Extracted image for figure '%s' (title: '%s', size: %d bytes) — description: NONE (GPT-4o Vision step not implemented)",
                                image_on_page.figure_id,
                                image_on_page.title or "no title",
                                len(image_on_page.bytes),
                            )

                page_text = page_text.replace("<!-- PageBreak -->", "")
                page_text = page_text.strip()
                if heading_map:
                    page_text = _patch_heading_levels(page_text, heading_map)
                yield Page(
                    page_num=page.page_number - 1,
                    offset=offset,
                    text=page_text,
                    images=page_images,
                    tables=page_tables,
                )
                offset += len(page_text)

    @staticmethod
    async def figure_to_image(doc: pymupdf.Document, figure: DocumentFigure) -> ImageOnPage:
        figure_title = figure.caption.content if figure.caption and figure.caption.content else ""
        figure_id = figure.id or f"fig_{uuid.uuid4().hex[:8]}"
        figure_filename = f"figure{figure_id.replace('.', '_')}.png"
        logger.info("Cropping figure %s with title '%s'", figure_id, figure_title)
        placeholder = f'<figure id="{figure_id}"></figure>'
        if not figure.bounding_regions:
            return ImageOnPage(
                bytes=b"",
                page_num=0,
                figure_id=figure_id,
                bbox=(0, 0, 0, 0),
                filename=figure_filename,
                title=figure_title,
                placeholder=placeholder,
                mime_type="image/png",
            )
        first_region = figure.bounding_regions[0]
        bounding_box = (
            first_region.polygon[0],
            first_region.polygon[1],
            first_region.polygon[4],
            first_region.polygon[5],
        )
        page_number = first_region.page_number
        cropped_img, bbox_pixels = DocumentAnalysisParser.crop_image_from_pdf_page(doc, page_number - 1, bounding_box)
        return ImageOnPage(
            bytes=cropped_img,
            page_num=page_number - 1,
            figure_id=figure_id,
            bbox=bbox_pixels,
            filename=figure_filename,
            title=figure_title,
            placeholder=placeholder,
            mime_type="image/png",
        )

    @staticmethod
    def table_to_html(table: DocumentTable):
        table_html = "<figure><table>"
        rows = [
            sorted([cell for cell in table.cells if cell.row_index == i], key=lambda cell: cell.column_index)
            for i in range(table.row_count)
        ]
        for row_cells in rows:
            table_html += "<tr>"
            for cell in row_cells:
                tag = "th" if (cell.kind == "columnHeader" or cell.kind == "rowHeader") else "td"
                cell_spans = ""
                if cell.column_span is not None and cell.column_span > 1:
                    cell_spans += f" colSpan={cell.column_span}"
                if cell.row_span is not None and cell.row_span > 1:
                    cell_spans += f" rowSpan={cell.row_span}"
                table_html += f"<{tag}{cell_spans}>{html.escape(cell.content)}</{tag}>"
            table_html += "</tr>"
        table_html += "</table></figure>"
        return table_html

    @staticmethod
    def crop_image_from_pdf_page(
        doc: pymupdf.Document, page_number: int, bbox_inches: tuple[float, float, float, float]
    ) -> tuple[bytes, tuple[float, float, float, float]]:
        """Crops a region from a given page in a PDF and returns it as an image."""
        bbox_dpi = 72
        x0, y0, x1, y1 = (round(x * bbox_dpi, 2) for x in bbox_inches)
        bbox_pixels = (x0, y0, x1, y1)
        rect = pymupdf.Rect(bbox_pixels)
        page_dpi = 300
        page = doc.load_page(page_number)
        pix = page.get_pixmap(matrix=pymupdf.Matrix(page_dpi / bbox_dpi, page_dpi / bbox_dpi), clip=rect)

        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        bytes_io = io.BytesIO()
        img.save(bytes_io, format="PNG")
        return bytes_io.getvalue(), bbox_pixels
